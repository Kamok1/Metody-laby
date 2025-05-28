from docx import Document
from docx.shared import Inches

# Create Word document
doc = Document()
doc.add_heading("MSID – Report (Part II)", level=1)
doc.add_paragraph("Author: Kamil Borusiak – 280447")

# --- DATA SPLIT ---
doc.add_heading("1. Data Split", level=2)
doc.add_paragraph("The dataset was divided into three parts:")
doc.add_paragraph("- Training set: 70%\n- Test set: 20%\n- Validation set: 10%")

# --- METRICS EXPLANATION ---
doc.add_heading("2. Evaluation Metrics", level=2)
doc.add_paragraph(
    "The following metrics were used to evaluate model performance:\n"
    "- Accuracy: The ratio of correctly predicted observations to the total observations.\n"
    "- Precision: The ratio of correctly predicted positive observations to the total predicted positives.\n"
    "- Recall: The ratio of correctly predicted positives to all actual positives.\n"
    "- F1-score: Harmonic mean of precision and recall.\n"
    "- AUC-ROC: Area under the Receiver Operating Characteristic curve, measuring the model’s ability to distinguish between classes.\n"
    "- MSE (Mean Squared Error), RMSE (Root MSE), MAE (Mean Absolute Error), R²: Standard regression performance metrics."
)

# --- CLASSIFICATION MODELS TABLES ---
doc.add_heading("3. Classification Model Performance", level=2)

models = {
    "Logistic Regression": [
        ["Set", "Class", "Precision", "Recall", "F1-score", "Support"],
        ["Test", "-1.0", "0.92", "0.86", "0.89", "217"],
        ["", "1.0", "0.91", "0.95", "0.93", "328"],
        ["", "Accuracy", "", "", "0.92", "545"],
        ["Validation", "-1.0", "0.93", "0.83", "0.87", "214"],
        ["", "1.0", "0.90", "0.96", "0.93", "330"],
        ["", "Accuracy", "", "", "0.91", "544"],
        ["Train", "-1.0", "0.94", "0.85", "0.89", "990"],
        ["", "1.0", "0.91", "0.97", "0.94", "1551"],
        ["", "Accuracy", "", "", "0.92", "2541"],
    ],
    "Decision Tree": [
        ["Set", "Class", "Precision", "Recall", "F1-score", "Support"],
        ["Test", "-1.0", "0.85", "0.83", "0.84", "217"],
        ["", "1.0", "0.89", "0.90", "0.90", "328"],
        ["", "Accuracy", "", "", "0.87", "545"],
        ["Validation", "-1.0", "0.82", "0.80", "0.81", "214"],
        ["", "1.0", "0.87", "0.89", "0.88", "330"],
        ["", "Accuracy", "", "", "0.85", "544"],
        ["Train", "-1.0", "1.00", "1.00", "1.00", "990"],
        ["", "1.0", "1.00", "1.00", "1.00", "1551"],
        ["", "Accuracy", "", "", "1.00", "2541"],
    ],
    "SVM": [
        ["Set", "Class", "Precision", "Recall", "F1-score", "Support"],
        ["Test", "-1.0", "0.94", "0.86", "0.90", "217"],
        ["", "1.0", "0.91", "0.97", "0.94", "328"],
        ["", "Accuracy", "", "", "0.92", "545"],
        ["Validation", "-1.0", "0.95", "0.82", "0.88", "214"],
        ["", "1.0", "0.89", "0.97", "0.93", "330"],
        ["", "Accuracy", "", "", "0.91", "544"],
        ["Train", "-1.0", "0.98", "0.86", "0.92", "990"],
        ["", "1.0", "0.92", "0.99", "0.95", "1551"],
        ["", "Accuracy", "", "", "0.94", "2541"],
    ]
}

for name, table_data in models.items():
    doc.add_heading(name, level=3)
    table = doc.add_table(rows=1, cols=len(table_data[0]))
    table.style = 'Table Grid'
    for i, header in enumerate(table_data[0]):
        table.cell(0, i).text = header
    for row_data in table_data[1:]:
        row = table.add_row().cells
        for i, item in enumerate(row_data):
            row[i].text = item

# --- CLOSED FORM LINEAR REGRESSION ---
doc.add_heading("4. Closed-Form Linear Regression", level=2)
table = doc.add_table(rows=4, cols=5)
table.style = 'Table Grid'
headers = ["Set", "MSE", "RMSE", "MAE", "R²"]
for i, h in enumerate(headers):
    table.cell(0, i).text = h
data = [
    ["Train", "0.2272", "0.4767", "0.2918", "0.9617"],
    ["Test", "0.2266", "0.4760", "0.3109", "0.9657"],
    ["Validation", "0.2396", "0.4895", "0.3143", "0.9641"],
]
for row_data in data:
    row = table.add_row().cells
    for i, item in enumerate(row_data):
        row[i].text = item

# --- LOGISTIC REGRESSION COMPARISON ---
doc.add_heading("5. Logistic Regression Summary", level=2)

for title, rows in {
    "Custom Logistic Regression": [
        ["Set", "Accuracy", "F1-score", "AUC-ROC"],
        ["Train", "0.921", "0.937", "0.961"],
        ["Test", "0.919", "0.935", "0.961"],
        ["Validation", "0.912", "0.930", "0.960"],
    ],
    "Scikit-learn Logistic Regression": [
        ["Set", "Accuracy", "F1-score", "AUC-ROC"],
        ["Train", "0.921", "0.937", "0.964"],
        ["Test", "0.916", "0.931", "0.962"],
        ["Validation", "0.906", "0.925", "0.960"],
    ]
}.items():
    doc.add_heading(title, level=3)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    for i, col in enumerate(rows[0]):
        table.cell(0, i).text = col
    for row_data in rows[1:]:
        row = table.add_row().cells
        for i, item in enumerate(row_data):
            row[i].text = item

# --- CPU vs GPU COMPARISON ---
doc.add_heading("6. CPU vs GPU Training Comparison", level=2)
doc.add_paragraph("Training time is longer on GPU due to the small model size and dataset. GPU has a higher overhead "
                  "from transferring data between CPU and GPU memory, and initializing CUDA kernels, which outweighs "
                  "the parallel processing benefits for this simple task.")

table = doc.add_table(rows=3, cols=8)
table.style = 'Table Grid'
headers = ["Device", "Train Time", "Test Acc", "Test F1", "Test AUC", "Val Acc", "Val F1", "Val AUC"]
for i, h in enumerate(headers):
    table.cell(0, i).text = h
table.add_row().cells[:] = ["CPU", "4.80 s", "0.934", "0.946", "0.960", "0.914", "0.931", "0.958"]
table.add_row().cells[:] = ["GPU", "6.71 s", "0.930", "0.943", "0.960", "0.914", "0.931", "0.958"]

# Save document
doc_path = "/MSID_Report_Kamil_Borusiak.docx"
doc.save(doc_path)
doc_path